# report_design_generator.py
# Run this script in VSCode with Python extension installed.
# First, install required library: pip install python-docx
# Then run: python report_design_generator.py
# Output: Report_Design.docx (Word file)

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor

# Create a new Document
doc = Document()

# Set default font for Thai support (user should have a Thai font like 'Angsana New' or 'TH Sarabun New')
style = doc.styles['Normal']
font = style.font
font.name = 'TH Sarabun New'  # Or 'Angsana New' if TH Sarabun not available
font.size = Pt(14)

# Title Page
title = doc.add_heading('รายงานการออกแบบหน้ารายงาน (Report Design) สำหรับระบบจัดการสมาชิกชมรม', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('ชื่อระบบ: ระบบจัดการสมาชิกชมรม (Club Member Management System)')
doc.add_paragraph('ชื่อเอกสาร: รายงานการออกแบบหน้ารายงาน (Report Design)')
doc.add_paragraph('วันที่: 16 กันยายน 2568')
doc.add_paragraph('ผู้จัดทำ: ระบบออกแบบโดย Grok (จากเอกสาร System Design)')
doc.add_paragraph('วัตถุประสงค์: ออกแบบรูปแบบรายงานตาม Output Design ในบทที่ 8 เพื่อให้ตรง Requirement จาก DFD และ Process Hierarchy')

doc.add_page_break()

# Table of Contents
doc.add_heading('สารบัญ', level=1)
toc = [
    '1. รายงานจำนวนสมาชิก .................................................. หน้า 2',
    '2. รายงานจำนวนกิจกรรม .................................................. หน้า 3',
    '3. รายงานการเข้าร่วมกิจกรรมของสมาชิก .............................. หน้า 4',
    '4. รายงานสรุปกิจกรรมทั้งหมด ......................................... หน้า 5',
    '**หมายเหตุทั่วไป:** รายงานทั้งหมดใช้ข้อมูลจาก D1 (สมาชิก), D2 (กิจกรรม), D3 (การเข้าร่วมกิจกรรม) สามารถพิมพ์หรือ export เป็น PDF ได้'
]
for item in toc:
    doc.add_paragraph(item)

doc.add_page_break()

# Report 1: Member Count
doc.add_heading('1. รายงานจำนวนสมาชิก (Member Count Report)', level=1)

# Header Table
header_table = doc.add_table(rows=1, cols=2)
header_table.style = 'Table Grid'
header_cell1 = header_table.cell(0, 0)
header_cell1.text = 'ชื่อระบบ: ระบบจัดการสมาชิกชมรม\nชื่อรายงาน: รายงานจำนวนสมาชิก\nวันที่พิมพ์: 16 กันยายน 2568\nผู้จัดทำ: เจ้าหน้าที่กิจกรรม'
header_cell2 = header_table.cell(0, 1)
header_cell2.text = '[พื้นที่สำหรับโลโก้ชมรม]'

# Body
doc.add_paragraph('สรุปจำนวนสมาชิกทั้งหมด: 150 คน (จากกระบวนการ 5.1 ใน DFD Level 1)')

body_table = doc.add_table(rows=5, cols=3)
body_table.style = 'Table Grid'
body_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Headers
hdr_cells = body_table.rows[0].cells
hdr_cells[0].text = 'สิทธิ์'
hdr_cells[1].text = 'จำนวน'
hdr_cells[2].text = 'เปอร์เซ็นต์'

# Data
rows_data = [
    ('Admin', '5', '3.33%'),
    ('Officer', '10', '6.67%'),
    ('Member', '135', '90%'),
    ('**รวม**', '**150**', '**100%**')
]
for i, row_data in enumerate(rows_data, 1):
    row_cells = body_table.rows[i].cells
    for j, val in enumerate(row_data):
        row_cells[j].text = val

# Footer
footer_table = doc.add_table(rows=1, cols=2)
footer_table.style = 'Table Grid'
footer_cell1 = footer_table.cell(0, 0)
footer_cell1.text = 'หมายเหตุ: ข้อมูลอัปเดตจากฐานข้อมูลล่าสุด (กระบวนการ 1: จัดการผู้ใช้งาน)\nผู้รับผิดชอบ: ประธานชมรม\nวันที่: 16/09/2568'
footer_cell2 = footer_table.cell(0, 1)
footer_cell2.text = '[พื้นที่สำหรับลายเซ็น]'

doc.add_page_break()

# Report 2: Event Count
doc.add_heading('2. รายงานจำนวนกิจกรรม (Event Count Report)', level=1)

# Header
header_table2 = doc.add_table(rows=1, cols=2)
header_table2.style = 'Table Grid'
header_cell1_2 = header_table2.cell(0, 0)
header_cell1_2.text = 'ชื่อระบบ: ระบบจัดการสมาชิกชมรม\nชื่อรายงาน: รายงานจำนวนกิจกรรม\nวันที่พิมพ์: 16 กันยายน 2568\nผู้จัดทำ: เจ้าหน้าที่กิจกรรม'
header_cell2_2 = header_table2.cell(0, 1)
header_cell2_2.text = '[พื้นที่สำหรับไอคอนปฏิทิน]'

doc.add_paragraph('สรุปจำนวนกิจกรรมทั้งหมด: 25 กิจกรรม (จากกระบวนการ 5.2 ใน DFD Level 1)')

body_table2 = doc.add_table(rows=5, cols=3)
body_table2.style = 'Table Grid'
body_table2.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells2 = body_table2.rows[0].cells
hdr_cells2[0].text = 'สถานะ'
hdr_cells2[1].text = 'จำนวน'
hdr_cells2[2].text = 'เปอร์เซ็นต์'

rows_data2 = [
    ('เปิดรับสมัคร', '10', '40%'),
    ('ปิดรับสมัคร', '12', '48%'),
    ('ยกเลิก', '3', '12%'),
    ('**รวม**', '**25**', '**100%**')
]
for i, row_data in enumerate(rows_data2, 1):
    row_cells = body_table2.rows[i].cells
    for j, val in enumerate(row_data):
        row_cells[j].text = val

doc.add_paragraph('[พื้นที่สำหรับกราฟ Pie Chart: แสดงสัดส่วนสถานะ (ใช้ Excel Insert Chart)]')

# Footer
footer_table2 = doc.add_table(rows=1, cols=2)
footer_table2.style = 'Table Grid'
footer_cell1_2 = footer_table2.cell(0, 0)
footer_cell1_2.text = 'หมายเหตุ: สถานะอัปเดตจากกระบวนการ 3.1 (จัดการสถานะกิจกรรม)\nผู้รับผิดชอบ: เจ้าหน้าที่กิจกรรม\nวันที่: 16/09/2568'
footer_cell2_2 = footer_table2.cell(0, 1)
footer_cell2_2.text = '[พื้นที่สำหรับลายเซ็น]'

doc.add_page_break()

# Report 3: Participation Report
doc.add_heading('3. รายงานการเข้าร่วมกิจกรรมของสมาชิก (Member Event Participation Report)', level=1)

header_table3 = doc.add_table(rows=1, cols=2)
header_table3.style = 'Table Grid'
header_cell1_3 = header_table3.cell(0, 0)
header_cell1_3.text = 'ชื่อระบบ: ระบบจัดการสมาชิกชมรม\nชื่อรายงาน: รายงานการเข้าร่วมกิจกรรมของสมาชิก\nวันที่พิมพ์: 16 กันยายน 2568\nผู้จัดทำ: ประธานชมรม'
header_cell2_3 = header_table3.cell(0, 1)
header_cell2_3.text = '[พื้นที่สำหรับไอคอนกลุ่มคน]'

doc.add_paragraph('สรุป: มี 120 การเข้าร่วมจาก 150 สมาชิก (จากกระบวนการ 5.3 ใน DFD Level 1 และ 4.4: ดูประวัติการเข้าร่วม)')

body_table3 = doc.add_table(rows=6, cols=5)  # Example with 5 rows data + header
body_table3.style = 'Table Grid'
body_table3.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells3 = body_table3.rows[0].cells
hdr_cells3[0].text = 'รหัสสมาชิก'
hdr_cells3[1].text = 'ชื่อ-นามสกุล'
hdr_cells3[2].text = 'ชื่อกิจกรรม'
hdr_cells3[3].text = 'วันที่เข้าร่วม'
hdr_cells3[4].text = 'สถานะ'

rows_data3 = [
    ('M001', 'นาย ก', 'กิจกรรมกีฬา', '01/09/2568', 'เสร็จสิ้น'),
    ('M002', 'นางสาว ข', 'กิจกรรมอบรม', '05/09/2568', 'เสร็จสิ้น'),
    ('M003', 'นาย ค', 'กิจกรรมกีฬา', '01/09/2568', 'เสร็จสิ้น'),
    ('**รวม**', '**120 รายการ**', '', '', '')
]
for i, row_data in enumerate(rows_data3, 1):
    row_cells = body_table3.rows[i].cells
    for j, val in enumerate(row_data):
        row_cells[j].text = val

doc.add_paragraph('... (เพิ่มแถวตามข้อมูลจริง, สูงสุด 50 แถวต่อหน้า)')

# Footer
footer_table3 = doc.add_table(rows=1, cols=2)
footer_table3.style = 'Table Grid'
footer_cell1_3 = footer_table3.cell(0, 0)
footer_cell1_3.text = 'หมายเหตุ: ข้อมูลจากกระบวนการ 3.2 (ลงทะเบียนเข้าร่วมกิจกรรม)\nผู้รับผิดชอบ: ประธานชมรม\nวันที่: 16/09/2568'
footer_cell2_3 = footer_table3.cell(0, 1)
footer_cell2_3.text = '[พื้นที่สำหรับลายเซ็น]'

doc.add_page_break()

# Report 4: Summary Report
doc.add_heading('4. รายงานสรุปกิจกรรมทั้งหมด (All Events Summary Report)', level=1)

header_table4 = doc.add_table(rows=1, cols=2)
header_table4.style = 'Table Grid'
header_cell1_4 = header_table4.cell(0, 0)
header_cell1_4.text = 'ชื่อระบบ: ระบบจัดการสมาชิกชมรม\nชื่อรายงาน: รายงานสรุปกิจกรรมทั้งหมด\nวันที่พิมพ์: 16 กันยายน 2568\nผู้จัดทำ: เจ้าหน้าที่กิจกรรม'
header_cell2_4 = header_table4.cell(0, 1)
header_cell2_4.text = '[พื้นที่สำหรับไอคอนกิจกรรม]'

doc.add_paragraph('สรุป: มี 25 กิจกรรม, ผู้เข้าร่วมรวม 300 คน (จากกระบวนการ 5.4 ใน DFD Level 1)')

body_table4 = doc.add_table(rows=5, cols=5)
body_table4.style = 'Table Grid'
body_table4.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells4 = body_table4.rows[0].cells
hdr_cells4[0].text = 'รหัสกิจกรรม'
hdr_cells4[1].text = 'ชื่อกิจกรรม'
hdr_cells4[2].text = 'วันที่'
hdr_cells4[3].text = 'สถานะ'
hdr_cells4[4].text = 'จำนวนผู้เข้าร่วม'

rows_data4 = [
    ('E001', 'กิจกรรมกีฬา', '01/09/2568', 'ปิด', '50'),
    ('E002', 'กิจกรรมอบรม', '05/09/2568', 'ปิด', '40'),
    ('E003', 'กิจกรรมสัมมนา', '10/09/2568', 'เปิดรับสมัคร', '30'),
    ('**รวม**', '**25 กิจกรรม**', '', '', '**300 คน**')
]
for i, row_data in enumerate(rows_data4, 1):
    row_cells = body_table4.rows[i].cells
    for j, val in enumerate(row_data):
        row_cells[j].text = val

doc.add_paragraph('[พื้นที่สำหรับ Bar Chart: ชื่อกิจกรรม vs จำนวนผู้เข้าร่วม (ใช้ Excel Insert Chart)]')

# Footer
footer_table4 = doc.add_table(rows=1, cols=2)
footer_table4.style = 'Table Grid'
footer_cell1_4 = footer_table4.cell(0, 0)
footer_cell1_4.text = 'หมายเหตุ: สรุปจากกระบวนการ 2 (จัดกิจกรรม) และ 3 (ควบคุมการเข้าร่วม)\nผู้รับผิดชอบ: เจ้าหน้าที่กิจกรรม\nวันที่: 16/09/2568'
footer_cell2_4 = footer_table4.cell(0, 1)
footer_cell2_4.text = '[พื้นที่สำหรับลายเซ็น]'

doc.add_page_break()

# Last Page: References
doc.add_heading('บรรทัดฐานและการอ้างอิง', level=1)
doc.add_paragraph('บรรทัดฐานการออกแบบ: ยึดตาม Output Design ใน "08 System Design.pdf" (หน้า 3-4): เริ่มจาก Requirement, สอดคล้อง Data Dictionary, รูปแบบ Report/กราฟิก.')
doc.add_paragraph('แหล่งข้อมูล: จาก DFD Level 0/1 ใน "Club-member-management-system-version-2.pdf" และ Process ใน "ระบบสมาชิก.pdf".')
doc.add_paragraph('เวอร์ชัน: 1.0 | อัปเดต: 16/09/2568')

# Save the document
doc.save('Report_Design.docx')
print("Generated Report_Design.docx successfully!")